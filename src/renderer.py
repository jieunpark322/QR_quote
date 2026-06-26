from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from .labels import DocumentLabels, load_labels
from .loader import load_clause, render_clause_body
from .models import Brand, QuoteDocument, ensure_totals


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _vcenter(cell) -> None:
    """셀의 수직 정렬을 가운데로 (XML 직접 조작 - LibreOffice 변환 호환)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:vAlign")):
        tc_pr.remove(old)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tc_pr.append(v)


def _zero_cell_lr_margin(cell) -> None:
    """셀의 좌우 padding(margin)을 0으로. 페이지 여백과 셀 내용을 정렬할 때 사용."""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is None:
        existing = OxmlElement("w:tcMar")
        tc_pr.append(existing)
    for direction in ("left", "right"):
        for old in existing.findall(qn(f"w:{direction}")):
            existing.remove(old)
        m = OxmlElement(f"w:{direction}")
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        existing.append(m)


def _set_cell_right_margin_cm(cell, cm_value: float) -> None:
    """셀의 오른쪽 padding(margin)을 cm 단위로 지정. 우측 셀 내용을 좌측으로 미는 용도."""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is None:
        existing = OxmlElement("w:tcMar")
        tc_pr.append(existing)
    for old in existing.findall(qn("w:right")):
        existing.remove(old)
    m = OxmlElement("w:right")
    m.set(qn("w:w"), str(int(cm_value * 567)))  # 1cm = 567 dxa
    m.set(qn("w:type"), "dxa")
    existing.append(m)


def _set_table_borders(table, color: str = "BFBFBF", size: int = 4) -> None:
    """표 전체에 격자선을 추가. size는 1/8 pt 단위 (4 = 0.5pt)."""
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color.lstrip("#"))
        borders.append(b)
    tbl_pr.append(borders)


def _force_fixed_column_widths(table, widths) -> None:
    """LibreOffice 변환 시 컬럼 너비를 보존하기 위해 tblLayout=fixed + tcW 강제 적용."""
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    # tblLayout 기존 것이 있으면 제거 후 fixed 로 설정
    existing = tbl_pr.find(qn("w:tblLayout"))
    if existing is not None:
        tbl_pr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed"),
    tbl_pr.append(layout)

    # 표 전체 너비를 명시적으로 설정 (twips 단위, 1cm = 567 twips)
    total_w_twips = int(sum(w.cm * 567 for w in widths))
    existing_w = tbl_pr.find(qn("w:tblW"))
    if existing_w is not None:
        tbl_pr.remove(existing_w)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(total_w_twips))
    tbl_pr.append(tblW)

    # 각 컬럼별 grid 너비 명시 (tblGrid)
    existing_grid = table._tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        table._tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w.cm * 567)))
        grid.append(gc)
    # tblGrid는 tblPr 다음에 와야 함
    table._tbl.insert(list(table._tbl).index(tbl_pr) + 1, grid)

    # 각 셀에 tcW 명시 (LibreOffice 가 cell.width 만으로는 무시할 수 있어 XML 강제)
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tcW_existing = tc_pr.find(qn("w:tcW"))
            if tcW_existing is not None:
                tc_pr.remove(tcW_existing)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(int(w.cm * 567)))
            tc_pr.append(tcW)


def _disable_auto_space_doc(doc) -> None:
    """문서 내 모든 단락의 autoSpaceDE/DN 을 0으로 설정.

    동아시아(한글)과 라틴/숫자가 이웃할 때 LibreOffice/Word 가 약 1/4 스페이스
    폭을 자동 삽입하는 기본 동작을 비활성화한다. "2026-06-26 까지" 같은 표현이
    스페이스 한 칸 더 들어간 것처럼 보이는 현상 제거.
    """
    body = doc.element.body
    for p in body.iter(qn("w:p")):
        p_pr = p.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            p.insert(0, p_pr)
        for tag in ("w:autoSpaceDE", "w:autoSpaceDN"):
            for old in p_pr.findall(qn(tag)):
                p_pr.remove(old)
            elem = OxmlElement(tag)
            elem.set(qn("w:val"), "0")
            p_pr.append(elem)


def _apply_font(run, font_name: str, *, size_pt: float | None = None,
                bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold:
        run.bold = True
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(doc, text: str, *, font: str, size_pt: float = 10,
                   bold: bool = False, alignment=None, color: RGBColor | None = None,
                   space_after_pt: float = 6, space_before_pt: float = 0):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.space_before = Pt(space_before_pt)
    run = p.add_run(text)
    _apply_font(run, font, size_pt=size_pt, bold=bold, color=color)
    return p


def _format_money(amount: float, currency: str = "KRW") -> str:
    if currency == "KRW":
        val = int(round(amount))
        if val < 0:
            return f"-₩{abs(val):,}"
        return f"₩{val:,}"
    return f"{amount:,.2f} {currency}"


# 할인(음수) 강조 색상 — 진한 빨강
DISCOUNT_COLOR = RGBColor(0xC0, 0x39, 0x2B)


def _render_logo(doc, brand: Brand, project_root: Path) -> None:
    """브랜드 로고가 있으면 문서 상단에 삽입. 없으면 조용히 건너뜀."""
    if not brand.branding.logo_path:
        return
    logo_full_path = project_root / "brands" / brand.brand_id / brand.branding.logo_path
    if not logo_full_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(logo_full_path), width=Cm(3.2))


def _render_header(doc, brand: Brand, document: QuoteDocument,
                   labels: DocumentLabels) -> None:
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)

    title_text = labels.quote.title if document.document_type == "quote" else labels.contract.title
    _add_paragraph(doc, title_text, font=font, size_pt=12, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, color=primary, space_after_pt=3)

    # 표 전체 폭은 USABLE_WIDTH 18.0cm (페이지 좌우 1.5cm 여백)
    # 우측 셀(7.0cm)을 inner 표(6.0cm)보다 1cm 넓게 잡아 inner 가 우측 정렬되며
    # 좌측에 1cm 빈 여백 확보 → 회색 박스가 좌측 회사정보에 딱 붙지 않도록
    LEFT_W = Cm(11.0)
    RIGHT_W = Cm(7.0)
    info_table = doc.add_table(rows=1, cols=2)
    info_table.autofit = False
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_table.columns[0].width = LEFT_W
    info_table.columns[1].width = RIGHT_W
    _force_fixed_column_widths(info_table, [LEFT_W, RIGHT_W])

    info_table.rows[0].height = Cm(2.0)
    info_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    left, right = info_table.rows[0].cells
    left.width = LEFT_W
    right.width = RIGHT_W
    _vcenter(left)
    _vcenter(right)
    # 좌측 셀 padding 제거 → "(주)소프트먼트" 가 페이지 여백과 정확히 정렬
    _zero_cell_lr_margin(left)

    p = left.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0)
    run = p.add_run(brand.company.name_ko)
    _apply_font(run, font, size_pt=11, bold=True, color=primary)

    info_lines = [
        f"사업자등록번호: {brand.company.registration_number}",
        f"대표자: {brand.company.ceo}",
    ]
    if brand.company.address:
        info_lines.append(brand.company.address)

    for line in info_lines:
        para = left.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.left_indent = Cm(0)
        r = para.add_run(line)
        _apply_font(r, font, size_pt=8.5)

    info_rows = [
        ("발행일", document.issued_date.isoformat()),
    ]
    if document.valid_until:
        info_rows.append(("유효기간", document.valid_until.isoformat() + " 까지"))
    if document.effective_date:
        info_rows.append(("계약시작일", document.effective_date.isoformat()))

    cp = brand.contact_person
    if cp:
        담당자_표기 = cp.name + (f" ({cp.title})" if cp.title else "")
        info_rows.append(("담당자", 담당자_표기))
        info_rows.append(("연락처", cp.phone or brand.company.phone or "-"))
        info_rows.append(("Email", cp.email or brand.company.email or "-"))
    else:
        info_rows.append(("연락처", brand.company.phone or "-"))
        info_rows.append(("Email", brand.company.email or "-"))

    inner = right.add_table(rows=len(info_rows), cols=2)
    inner.autofit = False
    inner.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # 우측 발행정보 표 (라벨 2.0cm + 값 4.0cm = 6.0cm)
    # VALUE_W 는 이메일 "shlee@softment.co.kr" 길이에 맞춰 잔여 공백 최소화
    LABEL_W = Cm(2.0)
    VALUE_W = Cm(4.0)
    inner.columns[0].width = LABEL_W
    inner.columns[1].width = VALUE_W
    for idx, (label, value) in enumerate(info_rows):
        row_obj = inner.rows[idx]
        # row 높이 명시 (LibreOffice가 vAlign 인식하도록).
        # AT_LEAST → 이메일 등 긴 텍스트는 줄바꿈하면서 행이 자동으로 늘어남 (잘림 방지)
        row_obj.height = Cm(0.4)
        row_obj.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        lc, vc = row_obj.cells
        # 셀 단위에도 너비 재지정 (LibreOffice가 columns[].width를 무시하는 경우 대비)
        lc.width = LABEL_W
        vc.width = VALUE_W
        _vcenter(lc)
        _vcenter(vc)
        _set_cell_bg(lc, "F0F2F8")
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        _apply_font(lr, font, size_pt=8.5, bold=True)
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(0)
        vp.paragraph_format.space_after = Pt(0)
        vr = vp.add_run(value)
        _apply_font(vr, font, size_pt=8.5)


def _render_counterparty(doc, brand: Brand, document: QuoteDocument,
                         labels: DocumentLabels) -> None:
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)

    _add_paragraph(doc, labels.quote.labels.counterparty_section,
                   font=font, size_pt=10, bold=True,
                   color=primary, space_before_pt=2, space_after_pt=2)

    cp = document.counterparty
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # 페이지 폭 가득 (들여쓰기 방지) — USABLE_WIDTH 18.0cm
    table.columns[0].width = Cm(18.0)
    _force_fixed_column_widths(table, [Cm(18.0)])
    cell = table.rows[0].cells[0]
    cell.width = Cm(18.0)
    _vcenter(cell)
    _set_cell_bg(cell, "FAFBFD")

    lines = [(cp.name, True, 11)]
    if cp.registration_number:
        lines.append((f"사업자등록번호: {cp.registration_number}", False, 8.5))
    if cp.address:
        lines.append((cp.address, False, 8.5))
    contact_bits = []
    if cp.contact_name:
        contact_bits.append(f"담당: {cp.contact_name}" + (f" ({cp.contact_title})" if cp.contact_title else ""))
    if cp.email:
        contact_bits.append(f"Email: {cp.email}")
    if contact_bits:
        lines.append(("  |  ".join(contact_bits), False, 8.5))

    first = cell.paragraphs[0]
    first.paragraph_format.space_after = Pt(1)
    fr = first.add_run(lines[0][0])
    _apply_font(fr, font, size_pt=lines[0][2], bold=lines[0][1])

    for text, bold, size in lines[1:]:
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        r = para.add_run(text)
        _apply_font(r, font, size_pt=size, bold=bold)

    _add_paragraph(doc, f"{labels.quote.labels.subject_prefix}: {document.subject}",
                   font=font, size_pt=9.5, bold=True,
                   space_before_pt=57, space_after_pt=2)


def _render_line_items(doc, brand: Brand, document: QuoteDocument,
                       labels: DocumentLabels) -> list | None:
    """품목 표 렌더링. 합계 표가 동일 컬럼 구조로 정렬할 수 있도록 widths 를 반환."""
    font = brand.branding.font_family
    primary_hex = brand.branding.colors.primary.lstrip("#")
    ql = labels.quote
    items = document.line_items

    if ql.labels.vat_separate_notice:
        vat_label_p = doc.add_paragraph()
        vat_label_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vat_label_p.paragraph_format.space_before = Pt(0)
        vat_label_p.paragraph_format.space_after = Pt(0)
        vat_run = vat_label_p.add_run(ql.labels.vat_separate_notice)
        _apply_font(vat_run, font, size_pt=8.5, bold=True,
                    color=RGBColor(0x99, 0x33, 0x33))

    th = ql.table_headers
    LEFT = WD_ALIGN_PARAGRAPH.LEFT
    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    RIGHT = WD_ALIGN_PARAGRAPH.RIGHT

    def _discount_text(i):
        if i.discount_amount and i.discount_amount > 0:
            # 할인 셀은 의미상 음수 — '-' 부호 없이 금액만 표시
            return _format_money(i.discount_amount, i.currency)
        if i.discount_rate:
            return f"{int(round(i.discount_rate * 100))}%"
        return ""

    # 컬럼 정의 — 새 순서: 항목·설명·단가·기간(횟수)·수량·할인·공급가·비고
    # 각 컬럼: (key, header, base_width_cm, align, value_getter, always_show, has_data_check)
    def _is_deferred(i) -> bool:
        return (getattr(i, "billing_type", None) or "") == "deferred_percent"

    all_columns = [
        ("name", th.name, 3.2, LEFT,
         lambda i: i.name or "", True, None),
        ("description", th.description, 3.2, LEFT,
         lambda i: i.description or "", False,
         lambda its: any((i.description or "").strip() for i in its)),
        ("unit_price", th.unit_price, 2.4, RIGHT,
         lambda i: ("-" if _is_deferred(i)
                    else _format_money(i.unit_price, i.currency)),
         True, None),
        ("period", th.period, 1.4, CENTER,
         lambda i: f"{i.period:g}" if (i.period is not None and i.period) else "",
         False,
         lambda its: any((i.period is not None and i.period and i.period != 0) for i in its)),
        ("qty", th.qty, 1.2, CENTER,
         lambda i: f"{i.qty:g}" if i.qty is not None else "", True, None),
        ("discount", "할인", 1.6, CENTER,
         _discount_text, False,
         lambda its: any(((i.discount_amount or 0) > 0) or i.discount_rate for i in its)),
        ("amount", th.amount, 2.8, RIGHT,
         lambda i: ("-" if _is_deferred(i)
                    else _format_money(i.amount, i.currency)),
         True, None),
        ("notes", th.notes, 3.2, LEFT,
         lambda i: i.notes or "", False,
         lambda its: any((i.notes or "").strip() for i in its)),
    ]

    # 활성 컬럼 필터링 (설명·기간·비고 중 모든 행이 비어있으면 숨김)
    active_cols = [
        c for c in all_columns
        if c[5] or (c[6] is not None and c[6](items))
    ]

    # ── 콘텐츠 한 줄 보장 우선 / 안 들어가면 폰트를 작게 줄여가며 시도 ──
    USABLE_WIDTH = 18.0      # 견적서 좌우 여백 1.5cm 가정
    PAD_CM = 0.45            # 셀 좌우 패딩
    MIN_SAFE_CM = 1.0        # 모든 컬럼 최소 폭

    # 폰트 후보 (큰 것부터) — 한글 한 글자 폭 (대략 폰트 pt × 0.032)
    CHAR_CM_BY_PT = {
        9.0:  0.285,
        8.5:  0.27,
        8.0:  0.255,
        7.5:  0.235,
        7.0:  0.22,
        6.5:  0.205,
        6.0:  0.19,
    }

    # 비고는 wrap 허용 — 긴 안내문구가 1~2줄에 들어가도록 충분히 큰 max
    NOTES_MAX = 5.0

    def _max_line_len(strings):
        m = 0
        for s in strings:
            if not s:
                continue
            for ln in str(s).split("\n"):
                m = max(m, len(ln))
        return m

    # 모든 컬럼의 콘텐츠 글자수 (헤더 포함) 미리 계산 — 폰트 시도 시 재사용
    content_lens = []
    for c in active_cols:
        _, header, _, _, getter, _, _ = c
        content_lens.append(_max_line_len(
            [str(getter(it)) for it in items] + [header]
        ))

    def _widths_for(char_cm: float) -> list[float]:
        """주어진 글자 폭(=폰트 크기) 으로 콘텐츠 한 줄 보장 너비 계산.
        비고는 NOTES_MAX 로 cap (wrap 허용)."""
        out = []
        for i, c in enumerate(active_cols):
            key = c[0]
            w = max(content_lens[i] * char_cm + PAD_CM, MIN_SAFE_CM)
            if key == "notes":
                w = min(w, NOTES_MAX)
            out.append(w)
        return out

    # 큰 폰트부터 시도 — 사용 가능 폭 안에 들어가면 그 폰트 채택
    chosen_font_pt = None
    raw_widths = None
    for font_pt in sorted(CHAR_CM_BY_PT.keys(), reverse=True):
        w_list = _widths_for(CHAR_CM_BY_PT[font_pt])
        if sum(w_list) <= USABLE_WIDTH:
            chosen_font_pt = font_pt
            raw_widths = w_list
            break

    if chosen_font_pt is None:
        # 가장 작은 폰트로도 안 들어감 — 설명을 먼저 양보(설명은 자체 \n 줄바꿈으로
        # 자연스럽게 wrap 됨), 그래도 부족하면 비고 양보
        chosen_font_pt = min(CHAR_CM_BY_PT.keys())
        raw_widths = _widths_for(CHAR_CM_BY_PT[chosen_font_pt])
        excess = sum(raw_widths) - USABLE_WIDTH
        # 1) 설명 먼저 양보 (최소 2.5cm 까지)
        for i, c in enumerate(active_cols):
            if c[0] == "description":
                give = min(excess, max(raw_widths[i] - 2.5, 0))
                raw_widths[i] -= give
                excess -= give
                break
        # 2) 그래도 초과면 비고 양보 (최소 2.5cm 까지)
        if excess > 0:
            for i, c in enumerate(active_cols):
                if c[0] == "notes":
                    give = min(excess, max(raw_widths[i] - 2.5, 0))
                    raw_widths[i] -= give
                    excess -= give
                    break
        # 3) 그래도 초과면 항목 양보 (최소 2.5cm 까지)
        if excess > 0:
            for i, c in enumerate(active_cols):
                if c[0] == "name":
                    raw_widths[i] = max(raw_widths[i] - excess, 2.5)
                    break
    else:
        # 남는 폭은 긴 컬럼(설명·항목·비고)에 가중 분배
        slack = USABLE_WIDTH - sum(raw_widths)
        if slack > 0.1:
            slack_weights_by_key = {
                "description": 3.0, "name": 1.5, "notes": 3.0,
            }
            weights = [slack_weights_by_key.get(c[0], 0.0) for c in active_cols]
            wsum = sum(weights)
            if wsum > 0:
                raw_widths = [w + slack * wt / wsum
                              for w, wt in zip(raw_widths, weights)]

    widths = [Cm(round(w, 2)) for w in raw_widths]
    headers = [c[1] for c in active_cols]

    table = doc.add_table(rows=1 + len(items), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(table, color="9FB0CC", size=4)
    # LibreOffice PDF 변환 시 컬럼 너비가 무시되지 않도록 layout fixed + tcW 강제
    _force_fixed_column_widths(table, widths)

    # 헤더 행
    header_row = table.rows[0]
    header_row.height = Cm(0.5)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for idx, (col_spec, width) in enumerate(zip(active_cols, widths)):
        cell = header_row.cells[idx]
        cell.width = width
        _vcenter(cell)
        _set_cell_bg(cell, primary_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(col_spec[1])
        _apply_font(r, font, size_pt=7, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # 폰트는 위 컬럼 너비 계산에서 결정된 chosen_font_pt 사용 — 한 줄 보장 우선
    cell_font_pt = chosen_font_pt
    # 행 높이는 항목 수 기반 (한 장 자동 채움)
    n_items = len(items)
    if n_items <= 4:
        row_h_cm = 1.8
    elif n_items <= 7:
        row_h_cm = 1.3
    elif n_items <= 11:
        row_h_cm = 1.0
    elif n_items <= 15:
        row_h_cm = 0.85
    elif n_items <= 20:
        row_h_cm = 0.7
    elif n_items <= 26:
        row_h_cm = 0.55
    else:
        row_h_cm = 0.5

    for r_idx, item in enumerate(items, start=1):
        row_obj = table.rows[r_idx]
        row_obj.height = Cm(row_h_cm)
        row_obj.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        is_discount = (item.amount or 0) < 0 or (item.unit_price or 0) < 0
        # 항목별 할인이 있는 일반 품목 행 — 할인 셀만 별도로 음영 강조
        item_has_disc = (
            (item.discount_amount and item.discount_amount > 0)
            or (item.discount_rate and item.discount_rate > 0)
        )
        for c_idx, (col_spec, width) in enumerate(zip(active_cols, widths)):
            col_key = col_spec[0]
            _, _, _, align, getter, _, _ = col_spec
            cell = row_obj.cells[c_idx]
            cell.width = width
            _vcenter(cell)
            # 음영 처리:
            #  - 할인 분류 행: 전체 셀 음영
            #  - 일반 품목 + 항목별 할인 입력: 그 행의 '할인' 셀만 음영
            if is_discount:
                _set_cell_bg(cell, "FDECEA")
            elif item_has_disc and col_key == "discount":
                _set_cell_bg(cell, "FDECEA")
            text_val = getter(item) or ""
            lines = str(text_val).split("\n") if text_val else [""]
            for ln_idx, ln in enumerate(lines):
                if ln_idx == 0:
                    p = cell.paragraphs[0]
                else:
                    p = cell.add_paragraph()
                p.alignment = align
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(ln)
                # 음영 처리한 할인 셀은 글자도 빨강+굵게
                cell_in_disc = is_discount or (
                    item_has_disc and col_key == "discount"
                )
                _apply_font(run, font, size_pt=cell_font_pt,
                            bold=cell_in_disc,
                            color=DISCOUNT_COLOR if cell_in_disc else None)

    return widths


def _render_totals(doc, brand: Brand, document: QuoteDocument,
                   labels: DocumentLabels,
                   item_table_widths: list | None = None) -> None:
    """합계 표를 품목 표와 동일한 컬럼 구조 + 가로 병합으로 만들어
    세로 선이 정확히 정렬되도록 한다."""
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)
    t = document.totals
    currency = t.currency

    ql = labels.quote.labels
    rows = []
    td_rate = getattr(document, "total_discount_rate", None) or 0
    if td_rate > 0:
        items_sum = sum(item.amount for item in document.line_items)
        td_value = items_sum - t.subtotal
        rows.append(("공급가액 (할인 전)", _format_money(items_sum, currency), False))
        rows.append((f"일괄 할인 ({int(round(td_rate * 100))}%)",
                     f"-{_format_money(td_value, currency)}", False))
    rows.extend([
        (ql.subtotal, _format_money(t.subtotal, currency), False),
        (f"{ql.vat} ({int(t.vat_rate * 100)}%)", _format_money(t.vat, currency), False),
        (ql.total, _format_money(t.total, currency), True),
    ])

    # 품목 표 widths 가 있으면 동일 컬럼 구조 + 라벨/값 가로 병합 → 세로 선 정렬
    if item_table_widths and len(item_table_widths) >= 2:
        n_cols = len(item_table_widths)
        # 값 영역은 마지막 2컬럼(공급가 + 비고) 가로 병합 → 큰 금액도 한 줄에 들어감
        value_col_count = 2 if n_cols >= 3 else 1
        label_col_count = n_cols - value_col_count
        table = doc.add_table(rows=len(rows), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_borders(table, color="9FB0CC", size=4)
        _force_fixed_column_widths(table, item_table_widths)
        for idx, (label, value, highlight) in enumerate(rows):
            row_obj = table.rows[idx]
            row_obj.height = Cm(0.75 if highlight else 0.65)
            row_obj.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            # 라벨 셀: col 0 ~ (label_col_count-1) 병합
            label_cell = row_obj.cells[0]
            for j in range(1, label_col_count):
                label_cell = label_cell.merge(row_obj.cells[j])
            # 값 셀: 라벨 다음 ~ 마지막 컬럼 병합
            value_cell = row_obj.cells[label_col_count]
            for j in range(label_col_count + 1, n_cols):
                value_cell = value_cell.merge(row_obj.cells[j])
            _vcenter(label_cell)
            _vcenter(value_cell)
            if highlight:
                _set_cell_bg(label_cell, brand.branding.colors.primary.lstrip("#"))
                _set_cell_bg(value_cell, brand.branding.colors.primary.lstrip("#"))
            # 라벨 (우측 정렬)
            lp = label_cell.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lp.paragraph_format.space_before = Pt(0)
            lp.paragraph_format.space_after = Pt(0)
            lr = lp.add_run(label)
            _apply_font(lr, font, size_pt=7, bold=highlight,
                        color=RGBColor(0xFF, 0xFF, 0xFF) if highlight else None)
            # 값 (우측 정렬)
            vp = value_cell.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            vp.paragraph_format.space_before = Pt(0)
            vp.paragraph_format.space_after = Pt(0)
            vr = vp.add_run(value)
            _apply_font(vr, font, size_pt=8 if highlight else 7, bold=highlight,
                        color=RGBColor(0xFF, 0xFF, 0xFF) if highlight else None)
        return

    # 폴백 — 품목 표 widths 가 없는 경우 (계약서 등) 기존 2컬럼 방식
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_table_borders(table, color="9FB0CC", size=4)
    for idx, (label, value, highlight) in enumerate(rows):
        row_obj = table.rows[idx]
        row_obj.height = Cm(0.75 if highlight else 0.65)
        row_obj.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        lc, vc = row_obj.cells
        lc.width = Cm(4)
        vc.width = Cm(4)
        _vcenter(lc)
        _vcenter(vc)
        if highlight:
            _set_cell_bg(lc, brand.branding.colors.primary.lstrip("#"))
            _set_cell_bg(vc, brand.branding.colors.primary.lstrip("#"))
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        _apply_font(lr, font, size_pt=10, bold=highlight,
                    color=RGBColor(0xFF, 0xFF, 0xFF) if highlight else None)
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vp.paragraph_format.space_before = Pt(0)
        vp.paragraph_format.space_after = Pt(0)
        vr = vp.add_run(value)
        _apply_font(vr, font, size_pt=11 if highlight else 10, bold=highlight,
                    color=RGBColor(0xFF, 0xFF, 0xFF) if highlight else None)


def _render_clauses(doc, brand: Brand, document: QuoteDocument, project_root: Path) -> None:
    if not document.clauses:
        return
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)

    _add_paragraph(doc, "약관 및 조건", font=font, size_pt=12, bold=True,
                   color=primary, space_after_pt=6)

    for number, clause_ref in enumerate(document.clauses, start=1):
        clause = load_clause(project_root, document.document_type, clause_ref.id)
        body = render_clause_body(clause, clause_ref, number)

        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(4)
        title_p.paragraph_format.space_after = Pt(2)
        tr = title_p.add_run(f"제 {number}조  {clause.title}")
        _apply_font(tr, font, size_pt=10.5, bold=True)

        body_p = doc.add_paragraph()
        body_p.paragraph_format.space_after = Pt(6)
        br = body_p.add_run(body.strip())
        _apply_font(br, font, size_pt=10)


def _render_etc_notice(doc, brand: Brand, document: QuoteDocument,
                       labels: DocumentLabels) -> None:
    """기타 안내 — 사용자가 textarea 에 입력한 notes 를 줄별 bullet 으로 표시.

    자동 안내(유효기간/입금계좌)는 webapp 의 textarea 기본값으로 채워지므로
    여기서는 추가하지 않음 (사용자가 본 그대로 PDF 에 반영).
    """
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)
    ql = labels.quote

    if not document.notes:
        return
    bullets = [ln.strip() for ln in document.notes.splitlines() if ln.strip()]
    if not bullets:
        return

    _add_paragraph(doc, ql.labels.etc_notice_section,
                   font=font, size_pt=10, bold=True,
                   color=primary, space_before_pt=10, space_after_pt=2)
    for text in bullets:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.left_indent = Cm(0.3)
        r = para.add_run(f"• {text}")
        _apply_font(r, font, size_pt=9.5)


def _render_signature(doc, brand: Brand, document: QuoteDocument,
                      labels: DocumentLabels) -> None:
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)

    _add_paragraph(doc, document.issued_date.strftime("%Y년 %m월 %d일"),
                   font=font, size_pt=11, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before_pt=28, space_after_pt=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(brand.company.name_ko)
    _apply_font(r, font, size_pt=14, bold=True, color=primary)

    if brand.footer_text:
        _add_paragraph(doc, brand.footer_text, font=font, size_pt=8.5,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       color=RGBColor(0x88, 0x88, 0x88),
                       space_before_pt=8, space_after_pt=0)


def _render_contract_title(doc, brand: Brand, document: QuoteDocument,
                           labels: DocumentLabels) -> None:
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)
    _add_paragraph(doc, labels.contract.title, font=font, size_pt=28, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, color=primary,
                   space_after_pt=6)
    if document.subject:
        _add_paragraph(doc, f"〈 {document.subject} 〉",
                       font=font, size_pt=13, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       space_after_pt=14)


def _render_contract_preamble(doc, brand: Brand, document: QuoteDocument,
                              labels: DocumentLabels) -> None:
    font = brand.branding.font_family
    preamble = labels.contract.preamble_template.format(
        supplier_name=brand.company.name_ko,
        counterparty_name=document.counterparty.name,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.line_spacing = 1.6
    r = p.add_run(preamble)
    _apply_font(r, font, size_pt=11)


def _render_contract_parties(doc, brand: Brand, document: QuoteDocument,
                             labels: DocumentLabels) -> None:
    font = brand.branding.font_family
    primary_hex = brand.branding.colors.primary.lstrip("#")
    cl = labels.contract.labels

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col, label in enumerate([cl.party_a, cl.party_b]):
        cell = table.rows[0].cells[col]
        _set_cell_bg(cell, primary_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        _apply_font(r, font, size_pt=11, bold=True,
                    color=RGBColor(0xFF, 0xFF, 0xFF))

    cp = document.counterparty
    parties = [
        [
            ("회사명", brand.company.name_ko),
            ("사업자등록번호", brand.company.registration_number),
            ("대표자", brand.company.ceo),
            ("주소", brand.company.address or "-"),
        ],
        [
            ("회사명", cp.name),
            ("사업자등록번호", cp.registration_number or "-"),
            ("대표자", cp.ceo or "-"),
            ("주소", cp.address or "-"),
        ],
    ]
    for col, lines in enumerate(parties):
        cell = table.rows[1].cells[col]
        _set_cell_bg(cell, "FAFBFD")
        for idx, (label, value) in enumerate(lines):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(f"{label}: ")
            _apply_font(r1, font, size_pt=9, bold=True)
            r2 = p.add_run(value or "-")
            _apply_font(r2, font, size_pt=9)
    doc.add_paragraph()


def _render_contract_overview(doc, brand: Brand, document: QuoteDocument,
                              labels: DocumentLabels) -> None:
    font = brand.branding.font_family
    cl = labels.contract.labels
    rows = []
    if document.effective_date:
        rows.append((cl.effective_date, document.effective_date.strftime("%Y년 %m월 %d일")))
    if document.contract_term:
        rows.append((cl.contract_term, document.contract_term))
    if document.totals:
        rows.append((cl.amount_pre_vat,
                     _format_money(document.totals.subtotal, document.totals.currency)))
        rows.append((cl.amount_with_vat,
                     _format_money(document.totals.total, document.totals.currency)))
    if not rows:
        return

    primary_hex = brand.branding.colors.primary.lstrip("#")
    _add_paragraph(doc, cl.overview_section, font=font, size_pt=12, bold=True,
                   color=_hex_to_rgb(brand.branding.colors.primary),
                   space_after_pt=4)

    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, value) in enumerate(rows):
        lc, vc = table.rows[idx].cells
        _set_cell_bg(lc, "F0F2F8")
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        _apply_font(lr, font, size_pt=10, bold=True)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        _apply_font(vr, font, size_pt=10)
    doc.add_paragraph()


def _render_contract_signature(doc, brand: Brand, document: QuoteDocument,
                               labels: DocumentLabels) -> None:
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)

    doc.add_paragraph()
    _add_paragraph(doc, document.issued_date.strftime("%Y년 %m월 %d일"),
                   font=font, size_pt=12, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_after_pt=20)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col, label in enumerate(["갑", "을"]):
        cell = table.rows[0].cells[col]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        _apply_font(r, font, size_pt=14, bold=True, color=primary)

    cp = document.counterparty
    for col, name in enumerate([brand.company.name_ko, cp.name]):
        cell = table.rows[1].cells[col]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        _apply_font(r, font, size_pt=13, bold=True)

    for col, addr in enumerate([brand.company.address or "", cp.address or ""]):
        cell = table.rows[2].cells[col]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(addr)
        _apply_font(r, font, size_pt=9)

    rep_a = f"{brand.signature.signer_title}  {brand.signature.signer_name}  (인)"
    cp_signer_title = "대표이사"
    cp_signer_name = cp.ceo or cp.contact_name or "_______________"
    rep_b = f"{cp_signer_title}  {cp_signer_name}  (인)"
    for col, rep in enumerate([rep_a, rep_b]):
        cell = table.rows[3].cells[col]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(rep)
        _apply_font(r, font, size_pt=11, bold=True)


def render_docx(brand: Brand, document: QuoteDocument, project_root: Path,
                output_path: Path) -> Path:
    labels = load_labels(project_root)
    # 누락된 totals 필드를 자동 계산 (공급가액 → 부가세 → 합계)
    ensure_totals(document, labels.quote.vat_rate)

    doc = Document()
    # 1페이지 보장을 위해 견적서는 여백을 빡빡하게 (단, 좌우는 1.5cm 이상으로 페이지 가장자리 여백 확보)
    is_quote = document.document_type == "quote"
    for section in doc.sections:
        section.top_margin = Cm(0.8 if is_quote else 1.2)
        section.bottom_margin = Cm(0.8 if is_quote else 1.2)
        section.left_margin = Cm(1.5 if is_quote else 1.8)
        section.right_margin = Cm(1.5 if is_quote else 1.8)

    _render_logo(doc, brand, project_root)

    if document.document_type == "contract":
        _render_contract_title(doc, brand, document, labels)
        _render_contract_preamble(doc, brand, document, labels)
        _render_contract_parties(doc, brand, document, labels)
        _render_contract_overview(doc, brand, document, labels)
        _render_clauses(doc, brand, document, project_root)
        _render_contract_signature(doc, brand, document, labels)
    else:
        _render_header(doc, brand, document, labels)
        _render_counterparty(doc, brand, document, labels)
        item_widths = _render_line_items(doc, brand, document, labels)
        _render_totals(doc, brand, document, labels,
                       item_table_widths=item_widths)
        _render_etc_notice(doc, brand, document, labels)
        _render_clauses(doc, brand, document, project_root)
        _render_signature(doc, brand, document, labels)

    # 한글-숫자 사이 자동 공백 제거 (모든 단락 일괄)
    _disable_auto_space_doc(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
