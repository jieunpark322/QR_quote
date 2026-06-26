"""멤버십 클라우드 견적서 DOCX 렌더러.

기존 quote 렌더러(평면 line_items)와 별도로 동작. 시나리오마다 페이지 분리,
계층 표(구분/분류/항목 + 종량제 하위), 자동 소계·총계.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .renderer import _force_fixed_column_widths, _set_cell_right_margin_cm
from .membership_models import (
    MembershipCategory,
    MembershipLineItem,
    MembershipQuoteDocument,
    MembershipScenario,
    MembershipSection,
    category_subtotal,
    scenario_grand_total_by_period,
    section_subtotals_by_period,
)
from .labels import load_labels
from .models import Brand


# ─── 색상/스타일 헬퍼 ────────────────────────────────────

def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _apply_font(run, font_name: str, *, size_pt: float | None = None,
                bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold:
        run.bold = True
    if color is not None:
        run.font.color.rgb = color


def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _vcenter(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:vAlign")):
        tc_pr.remove(old)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tc_pr.append(v)


def _set_table_borders(table, color: str = "9FB0CC", size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color.lstrip("#"))
        borders.append(b)
    tbl_pr.append(borders)


def _merge_vertical(table, col: int, start_row: int, end_row: int) -> None:
    """col(0-indexed) 컬럼의 start_row..end_row 를 수직 병합."""
    if end_row <= start_row:
        return
    first = table.cell(start_row, col)
    last = table.cell(end_row, col)
    first.merge(last)


def _add_paragraph(doc, text: str, *, font: str, size_pt: float = 10,
                   bold: bool = False, alignment=None, color: RGBColor | None = None,
                   space_after_pt: float = 4, space_before_pt: float = 0):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.space_before = Pt(space_before_pt)
    run = p.add_run(text)
    _apply_font(run, font, size_pt=size_pt, bold=bold, color=color)
    return p


def _format_won(amount: float) -> str:
    return f"{int(round(amount)):,}"


def _money(amount: float) -> str:
    """부호 우선 통화 표기 (-₩1,000,000 / ₩1,000,000)."""
    val = int(round(amount))
    if val < 0:
        return f"-₩{abs(val):,}"
    return f"₩{val:,}"


# 할인(음수) 강조 — 진한 빨강 + 연한 빨강 배경
DISCOUNT_COLOR = RGBColor(0xC0, 0x39, 0x2B)
DISCOUNT_BG = "FDECEA"


def _compute_section_widths(section: MembershipSection,
                            usable_cm: float = 19.5) -> list:
    """콘텐츠 길이 + 중요도 기반으로 컬럼 너비를 유동 산출.

    [구분, 분류, 상세 구분, 기간, 단가, 할인, 금액, 비고]
    """
    cats = section.categories
    items = [it for cat in cats for it in cat.items]

    # 각 컬럼의 최대 콘텐츠 글자 수 (한 줄 기준, 가장 긴 줄)
    def _max_line_len(strings):
        m = 0
        for s in strings:
            if not s:
                continue
            for ln in str(s).split("\n"):
                m = max(m, len(ln))
        return m

    # 헤더는 기본 한국어 2~3자 (구분, 분류, 기간 등)
    sec_len = _max_line_len([section.name, "구분"])
    # 분류 컬럼은 '예상 총 금액' 섹션 합계 라벨이 한 줄에 들어가야 함
    cat_len = _max_line_len([c.name for c in cats] + ["분류", "예상 총 금액"])

    # 상세 구분: name + name_detail + sub_items 줄들 최대치
    detail_strings = []
    for it in items:
        detail_strings.append(it.name)
        if it.name_detail:
            detail_strings.append(it.name_detail)
        if it.sub_items:
            # 종량제 줄 합치기 패턴 그대로
            sub_lines = []
            for s in it.sub_items:
                if s.label:
                    sub_lines.append(f"· {s.label}  {s.spec}")
                elif sub_lines:
                    sub_lines[-1] += f"  {s.spec}"
                else:
                    sub_lines.append(f"   {s.spec}")
            detail_strings.extend(sub_lines)
    detail_len = max(_max_line_len(detail_strings + ["상세 구분"]), 8)

    period_len = _max_line_len([it.billing_period for it in items] + ["기간"])

    unit_price_strings = []
    for it in items:
        if it.unit_price is not None:
            unit_price_strings.append(f"₩{int(it.unit_price):,}")
        elif it.unit_price_text:
            unit_price_strings.append(it.unit_price_text)
    unit_price_len = _max_line_len(unit_price_strings + ["단가"])

    amount_strings = []
    for it in items:
        if it.amount_text:
            amount_strings.append(it.amount_text)
        elif it.effective_amount() is not None:
            amount_strings.append(f"₩{int(it.effective_amount()):,}")
        else:
            amount_strings.append("-")
    # 분류 합계도 금액 컬럼에 표시
    for cat in cats:
        st = category_subtotal(cat)
        if st:
            amount_strings.append(f"₩{int(st):,}")
    amount_len = _max_line_len(amount_strings + ["금액"])

    notes_len = _max_line_len([it.notes or "" for it in items] + ["비고"])

    # ── 콘텐츠 한 줄 보장 우선 — 안 들어가면 폰트 작게 시도 ──
    PAD_CM = 0.4
    MIN_SAFE_CM = 1.0
    NOTES_MAX = 2.8
    DETAIL_MAX = 7.5

    CHAR_CM_BY_PT = {
        8.5: 0.27, 8.0: 0.255, 7.5: 0.235,
        7.0: 0.22, 6.5: 0.205, 6.0: 0.19,
    }

    content_lens = [sec_len, cat_len, detail_len, period_len,
                    unit_price_len, 4, amount_len, notes_len]
    headers_for_min = ["구분", "분류", "상세 구분", "기간", "단가",
                       "할인", "금액", "비고"]

    def _widths_for(char_cm: float) -> list[float]:
        out = []
        for i, h in enumerate(headers_for_min):
            w = max(
                content_lens[i] * char_cm + PAD_CM,
                len(h) * char_cm + PAD_CM,
                MIN_SAFE_CM,
            )
            if i == 2:
                w = min(w, DETAIL_MAX)
            elif i == 7:
                w = min(w, NOTES_MAX)
            out.append(w)
        return out

    # 큰 폰트부터 시도 — 사용 가능 폭에 들어가는 첫 폰트 채택
    chosen_font_pt = None
    constrained = None
    for font_pt in sorted(CHAR_CM_BY_PT.keys(), reverse=True):
        w_list = _widths_for(CHAR_CM_BY_PT[font_pt])
        if sum(w_list) <= usable_cm:
            chosen_font_pt = font_pt
            constrained = w_list
            break
    if chosen_font_pt is None:
        chosen_font_pt = min(CHAR_CM_BY_PT.keys())
        constrained = _widths_for(CHAR_CM_BY_PT[chosen_font_pt])

    # 가장 작은 폰트로도 폭 초과면 비고·상세에서 추가 양보 (한 줄 보장 깨질 수 있음)
    total = sum(constrained)
    if total > usable_cm:
        excess = total - usable_cm
        # 축소 우선순위 가중치 (비고 > 상세)
        shrink_weights = [0.0, 0.0, 2.0, 0.0, 0.0, 0.0, 0.0, 3.0]
        sw_sum = sum(shrink_weights)
        if sw_sum > 0:
            constrained = [
                w - excess * shrink_weights[i] / sw_sum
                for i, w in enumerate(constrained)
            ]
            # 최소 폭 클램프 (모든 컬럼 0.5cm 이상)
            constrained = [max(w, 0.5) for w in constrained]
            # 그래도 초과면 비고/상세에서 추가 양보
            total2 = sum(constrained)
            if total2 > usable_cm:
                constrained[2] = max(constrained[2] - (total2 - usable_cm), 2.5)
    else:
        # 남는 공간은 콘텐츠 폭이 큰 컬럼(상세구분·비고·구분)에 가중 분배
        slack = usable_cm - total
        slack_weights = [0.5, 0.3, 4.0, 0.0, 0.5, 0.0, 0.3, 1.5]
        sw_sum = sum(slack_weights)
        if sw_sum > 0:
            constrained = [
                w + slack * sw / sw_sum
                for w, sw in zip(constrained, slack_weights)
            ]

    return [Cm(round(w, 2)) for w in constrained]


# ─── 렌더링 함수들 ────────────────────────────────────────

def _render_logo(doc, brand: Brand, project_root: Path) -> None:
    if not brand.branding.logo_path:
        return
    logo_path = project_root / "brands" / brand.brand_id / brand.branding.logo_path
    if not logo_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(logo_path), width=Cm(2.4))


def _render_title(doc, document: MembershipQuoteDocument,
                  scenario: MembershipScenario, brand: Brand) -> None:
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)
    suffix = f" ({scenario.name})" if (scenario.name and scenario.name != "기본") else ""
    _add_paragraph(
        doc, f"{document.title}{suffix}",
        font=font, size_pt=11, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        color=primary, space_after_pt=1,
    )
    if scenario.subject:
        _add_paragraph(
            doc, scenario.subject,
            font=font, size_pt=8.5,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after_pt=6,
        )


def _render_header_qr_style(doc, document: MembershipQuoteDocument,
                            brand: Brand) -> None:
    """QR 견적서와 동일한 헤더 — 좌: 회사정보, 우: 발급정보 (발행일/담당자)."""
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)

    # 제목
    _add_paragraph(doc, document.title, font=font, size_pt=12, bold=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, color=primary,
                   space_after_pt=3)

    # 좌:회사정보 (넓게) | 우:발급정보 (컴팩트) — 합 = USABLE_WIDTH 19.5cm
    LEFT_W = Cm(12.5)
    RIGHT_W = Cm(7.0)
    info_table = doc.add_table(rows=1, cols=2)
    info_table.autofit = False
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_table.columns[0].width = LEFT_W
    info_table.columns[1].width = RIGHT_W
    _force_fixed_column_widths(info_table, [LEFT_W, RIGHT_W])
    info_table.rows[0].height = Cm(3.0)
    info_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    left, right = info_table.rows[0].cells
    left.width = LEFT_W
    right.width = RIGHT_W
    _vcenter(left)
    _vcenter(right)
    # 우측 발행정보 박스를 오른쪽 여백으로부터 1cm 안쪽으로 밀어 배치
    _set_cell_right_margin_cm(right, 1.0)

    # 좌측: 회사명 + 회사 정보
    p = left.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0)
    run = p.add_run(brand.company.name_ko)
    _apply_font(run, font, size_pt=11, bold=True, color=primary)
    info_lines = [
        f"사업자등록번호: {brand.company.registration_number}",
        f"대표자: {brand.company.ceo}",
    ]
    if brand.company.address:
        info_lines.append(brand.company.address)
    for line in info_lines:
        para = left.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.left_indent = Cm(0)
        r = para.add_run(line)
        _apply_font(r, font, size_pt=8.5)

    # 우측: 발행일 + 담당자 정보 (라벨 + 값 2컬럼 inner 표)
    info_rows = [("발행일", document.issued_date.isoformat())]
    cpn = brand.contact_person
    if cpn:
        label = cpn.name + (f" ({cpn.title})" if cpn.title else "")
        info_rows.append(("담당자", label))
        info_rows.append(("연락처", cpn.phone or brand.company.phone or "-"))
        info_rows.append(("Email", cpn.email or brand.company.email or "-"))
    else:
        info_rows.append(("연락처", brand.company.phone or "-"))
        info_rows.append(("Email", brand.company.email or "-"))

    inner = right.add_table(rows=len(info_rows), cols=2)
    inner.autofit = False
    inner.alignment = WD_TABLE_ALIGNMENT.RIGHT
    LABEL_W = Cm(1.5)
    VALUE_W = Cm(4.8)
    inner.columns[0].width = LABEL_W
    inner.columns[1].width = VALUE_W
    for idx, (label, value) in enumerate(info_rows):
        row_obj = inner.rows[idx]
        row_obj.height = Cm(0.6)
        row_obj.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        lc, vc = row_obj.cells
        lc.width = LABEL_W
        vc.width = VALUE_W
        _vcenter(lc)
        _vcenter(vc)
        _set_cell_bg(lc, "F0F2F8")
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        _apply_font(lr, font, size_pt=8.5, bold=True)
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(0)
        vp.paragraph_format.space_after = Pt(0)
        vr = vp.add_run(value)
        _apply_font(vr, font, size_pt=8.5)


def _render_counterparty_qr_style(doc, document: MembershipQuoteDocument,
                                   brand: Brand) -> None:
    """QR 견적서 스타일 — 제휴사(수신처) 박스 + 건명."""
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)

    _add_paragraph(doc, "수 신",
                   font=font, size_pt=10, bold=True,
                   color=primary, space_before_pt=14, space_after_pt=2)

    cp = document.counterparty
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(19.5)
    _force_fixed_column_widths(table, [Cm(19.5)])
    cell = table.rows[0].cells[0]
    cell.width = Cm(19.5)
    _vcenter(cell)
    _set_cell_bg(cell, "FAFBFD")

    lines = [(cp.name, True, 11)]
    if cp.address:
        lines.append((f"주소: {cp.address}", False, 8.5))
    if cp.ceo:
        lines.append((f"대표이사: {cp.ceo}", False, 8.5))
    contact_bits = []
    if cp.contact:
        contact_bits.append(f"담당: {cp.contact}")
    if cp.contact_phone:
        contact_bits.append(f"연락처: {cp.contact_phone}")
    if cp.contact_email:
        contact_bits.append(f"Email: {cp.contact_email}")
    if contact_bits:
        lines.append(("  |  ".join(contact_bits), False, 8.5))

    first = cell.paragraphs[0]
    first.paragraph_format.space_after = Pt(1)
    fr = first.add_run(lines[0][0])
    _apply_font(fr, font, size_pt=lines[0][2], bold=lines[0][1])
    for text, bold, size in lines[1:]:
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        r = para.add_run(text)
        _apply_font(r, font, size_pt=size, bold=bold)


def _render_signature_qr_style(doc, document: MembershipQuoteDocument,
                                brand: Brand) -> None:
    """QR 견적서와 동일한 서명 — 날짜 + 회사명 가운데 정렬."""
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)

    _add_paragraph(doc, document.issued_date.strftime("%Y년 %m월 %d일"),
                   font=font, size_pt=11,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before_pt=24, space_after_pt=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(brand.company.name_ko)
    _apply_font(r, font, size_pt=14, bold=True, color=primary)
    if brand.footer_text:
        _add_paragraph(doc, brand.footer_text, font=font, size_pt=8.5,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       color=RGBColor(0x88, 0x88, 0x88),
                       space_before_pt=8, space_after_pt=0)


def _render_parties(doc, document: MembershipQuoteDocument, brand: Brand) -> None:
    """[Legacy] 양측 발행 정보 (제휴사 | 회사). 호환용 유지 — 새 코드는 _render_header_qr_style 사용."""
    font = brand.branding.font_family
    primary_hex = brand.branding.colors.primary.lstrip("#")

    cp = document.counterparty
    sup = document.supplier
    if sup is None:
        # brand 정보로 자동 채움
        from .membership_models import MembershipParty
        cpn = brand.contact_person
        contact_label = None
        if cpn:
            contact_label = (
                f"{cpn.name} ({cpn.title})" if cpn.title else cpn.name
            )
        sup = MembershipParty(
            label="회사",
            name=brand.company.name_ko,
            address=brand.company.address,
            ceo=brand.company.ceo,
            contact=contact_label,
            contact_phone=(cpn.phone if cpn else None),
            contact_email=(cpn.email if cpn else None),
        )

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, color="BFBFBF", size=4)

    # 헤더 (라벨)
    for col, party in enumerate([cp, sup]):
        c = table.rows[0].cells[col]
        _vcenter(c)
        _set_cell_bg(c, primary_hex)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(party.label)
        _apply_font(r, font, size_pt=8.5, bold=True,
                    color=RGBColor(0xFF, 0xFF, 0xFF))
    table.rows[0].height = Cm(0.45)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # 정보
    for col, party in enumerate([cp, sup]):
        c = table.rows[1].cells[col]
        _vcenter(c)
        # 첫 줄: 회사명 (굵게)
        first = c.paragraphs[0]
        first.paragraph_format.space_before = Pt(0)
        first.paragraph_format.space_after = Pt(0)
        fr = first.add_run(party.name)
        _apply_font(fr, font, size_pt=9, bold=True)
        # 보조 정보
        info_lines = []
        if party.address:
            info_lines.append(f"주소 : {party.address}")
        if party.ceo:
            info_lines.append(f"대표이사 : {party.ceo}")
        if party.contact:
            info_lines.append(f"담당자 : {party.contact}")
        if party.contact_phone:
            info_lines.append(f"연락처 : {party.contact_phone}")
        if party.contact_email:
            info_lines.append(f"이메일 : {party.contact_email}")
        for line in info_lines:
            p = c.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            _apply_font(r, font, size_pt=7.5)


def _render_table_header_row(table, row_idx: int, headers: list[str], widths: list[Cm],
                             font: str, primary_hex: str) -> None:
    row = table.rows[row_idx]
    row.height = Cm(0.5)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for c_idx, (h, w) in enumerate(zip(headers, widths)):
        c = row.cells[c_idx]
        c.width = w
        _vcenter(c)
        _set_cell_bg(c, primary_hex)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        _apply_font(r, font, size_pt=7, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))


def _fill_cell(cell, text: str, *, font: str, size_pt: float = 9,
               bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT,
               color: RGBColor | None = None, bg: str | None = None) -> None:
    _vcenter(cell)
    if bg:
        _set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    lines = text.split("\n")
    for i, ln in enumerate(lines):
        if i == 0:
            r = p.add_run(ln)
        else:
            p2 = cell.add_paragraph()
            p2.alignment = align
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            r = p2.add_run(ln)
        _apply_font(r, font, size_pt=size_pt, bold=bold, color=color)


def _render_item_row(table, row_idx: int, item: MembershipLineItem, *,
                     widths: list[Cm], font: str) -> None:
    row = table.rows[row_idx]
    row.height = Cm(0.8)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    cells = row.cells

    # 상세 구분 (이름 + 보조 + 종량제 하위)
    name_text = item.name
    if item.name_detail:
        name_text += "\n" + item.name_detail
    if item.sub_items:
        # 종량제: 본행은 비우고 하위만 보여주거나, 본행 + 줄바꿈으로 하위 나열
        sub_lines = []
        for s in item.sub_items:
            if s.label:
                sub_lines.append(f"· {s.label}  {s.spec}")
            else:
                sub_lines.append(f"   {s.spec}")
        name_text += "\n" + "\n".join(sub_lines)
    _fill_cell(cells[0], name_text, font=font, size_pt=8.5,
               align=WD_ALIGN_PARAGRAPH.LEFT)

    # 기간
    _fill_cell(cells[1], item.billing_period or "-",
               font=font, size_pt=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 단가 (숫자 또는 텍스트)
    if item.unit_price is not None:
        up_text = f"₩{_format_won(item.unit_price)}"
    elif item.unit_price_text:
        up_text = item.unit_price_text
    else:
        up_text = "-"
    _fill_cell(cells[2], up_text, font=font, size_pt=8.5,
               align=WD_ALIGN_PARAGRAPH.RIGHT)

    # 할인율
    if item.discount_rate:
        dr_text = f"{int(item.discount_rate * 100)}%"
    else:
        dr_text = "-"
    _fill_cell(cells[3], dr_text, font=font, size_pt=8.5,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    # 금액
    eff = item.effective_amount()
    if item.amount_text:
        amt_text = item.amount_text
    elif eff is not None:
        amt_text = f"₩{_format_won(eff)}"
    else:
        amt_text = "-"
    _fill_cell(cells[4], amt_text, font=font, size_pt=8.5, bold=True,
               align=WD_ALIGN_PARAGRAPH.RIGHT)

    # 비고
    _fill_cell(cells[5], item.notes or "-",
               font=font, size_pt=8.5, align=WD_ALIGN_PARAGRAPH.LEFT)


def _render_section_table(doc, section: MembershipSection, brand: Brand) -> None:
    """한 구분의 표를 렌더링. 구분 컬럼은 좌측에 병합."""
    font = brand.branding.font_family
    primary_hex = brand.branding.colors.primary.lstrip("#")
    light_bg = "F0F2F8"

    # 표 컬럼: [구분 | 분류 | 상세구분 | 기간 | 단가 | 할인 | 금액 | 비고] 8개
    # 첫 컬럼(구분)은 첫 행에만 표시하고 나머지는 수직 병합

    # 행 수 계산: 헤더(1) + 각 분류별 [항목들 + (옵션) 합계 1행] + 섹션 합계 1행
    body_rows = 0
    for cat in section.categories:
        body_rows += len(cat.items)
        if cat.show_subtotal:
            body_rows += 1
    if section.show_section_total:
        body_rows += 1
    total_rows = 1 + body_rows  # +1 헤더

    cols_n = 8
    table = doc.add_table(rows=total_rows, cols=cols_n)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(table, color="BFBFBF", size=4)

    headers = ["구분", "분류", "상세 구분", "기간", "단가", "할인", "금액", "비고"]
    # 콘텐츠 길이 + 중요도 기반 유동 컬럼 너비
    widths = _compute_section_widths(section)
    # LibreOffice PDF 변환 시 컬럼 너비가 무시되지 않도록 layout fixed + tcW 강제
    _force_fixed_column_widths(table, widths)
    _render_table_header_row(table, 0, headers, widths, font, primary_hex)

    # 본문 채우기
    r = 1
    section_first_row = r  # 구분 컬럼 병합 시작점

    for cat in section.categories:
        cat_first_row = r  # 분류 컬럼 병합 시작점

        cat_is_discount = (cat.name or "").strip() == "할인"
        for item in cat.items:
            # 컬럼 0,1 은 분류 row에서만 텍스트 (병합 후 빈 값)
            _fill_cell(table.cell(r, 0), "", font=font,
                       bg=DISCOUNT_BG if cat_is_discount else None)
            _fill_cell(table.cell(r, 1), "", font=font,
                       bg=DISCOUNT_BG if cat_is_discount else None)
            # 컬럼 2~7: 항목 데이터
            _render_item_row_in_table(table, r, item, widths, font,
                                      is_discount=cat_is_discount)
            r += 1

        # 분류 컬럼 텍스트 (병합 후의 첫 행 = cat_first_row)
        _fill_cell(table.cell(cat_first_row, 1), cat.name,
                   font=font, size_pt=8, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   bg=DISCOUNT_BG if cat_is_discount else light_bg,
                   color=DISCOUNT_COLOR if cat_is_discount else None)
        # 분류 컬럼 수직 병합 (옵션 분류는 합계 행 없을 수 있음)
        cat_last_row = r - 1  # 마지막 항목 row
        if cat_first_row != cat_last_row:
            _merge_vertical(table, 1, cat_first_row, cat_last_row)

        # 분류 합계 행 (col 2~5 병합으로 라벨 넓게, col 6 금액, col 7 비고)
        if cat.show_subtotal:
            subtotal = category_subtotal(cat)
            label = cat.subtotal_label or f"{cat.name} 합계"
            subtotal_row = table.rows[r]
            subtotal_row.height = Cm(0.45)
            subtotal_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            sub_bg = DISCOUNT_BG if cat_is_discount else light_bg
            sub_color = DISCOUNT_COLOR if cat_is_discount else None
            _fill_cell(table.cell(r, 0), "", font=font,
                       bg=DISCOUNT_BG if cat_is_discount else None)
            _fill_cell(table.cell(r, 1), "", font=font,
                       bg=DISCOUNT_BG if cat_is_discount else None)
            label_cell = table.cell(r, 2).merge(table.cell(r, 5))
            _fill_cell(label_cell, label, font=font, size_pt=7, bold=True,
                       align=WD_ALIGN_PARAGRAPH.RIGHT,
                       bg=sub_bg, color=sub_color)
            _fill_cell(table.cell(r, 6),
                       _money(subtotal) if subtotal else "-",
                       font=font, size_pt=7, bold=True,
                       align=WD_ALIGN_PARAGRAPH.RIGHT,
                       bg=sub_bg, color=sub_color)
            _fill_cell(table.cell(r, 7), "-", font=font, size_pt=7,
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       bg=sub_bg, color=sub_color)
            r += 1

    # 구분 컬럼 텍스트 + 수직 병합 (섹션 합계 행 직전까지)
    _fill_cell(table.cell(section_first_row, 0), section.name,
               font=font, size_pt=8.5, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg=light_bg)
    section_body_last = r - 1
    if section_first_row != section_body_last:
        _merge_vertical(table, 0, section_first_row, section_body_last)

    # 섹션 예상 총 금액 행 (col 2~7 가로 병합으로 한 줄에 표시)
    if section.show_section_total:
        by_period = section_subtotals_by_period(section)
        parts = [f"{period}: {_money(amt)}" for period, amt in by_period.items()]
        total_text = "  /  ".join(parts) if parts else "-"

        section_total_row = table.rows[r]
        section_total_row.height = Cm(0.5)
        section_total_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        _fill_cell(table.cell(r, 0), "", font=font, bg=light_bg)
        _fill_cell(table.cell(r, 1), "예상 총 금액",
                   font=font, size_pt=7, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   bg=brand.branding.colors.primary.lstrip("#"),
                   color=RGBColor(0xFF, 0xFF, 0xFF))
        # col 2~7 가로 병합해서 넓은 한 셀로
        merged_cell = table.cell(r, 2).merge(table.cell(r, 7))
        _fill_cell(merged_cell, total_text,
                   font=font, size_pt=7.5, bold=True,
                   align=WD_ALIGN_PARAGRAPH.LEFT, bg=light_bg)
        r += 1


def _render_item_row_in_table(table, row_idx: int, item: MembershipLineItem,
                              widths: list[Cm], font: str,
                              is_discount: bool = False) -> None:
    """8컬럼 표에서 상세구분 시작 컬럼(idx=2)부터 채움.
    is_discount=True 이면 할인 행으로 강조 (빨간색 글자 + 배경).
    """
    row = table.rows[row_idx]
    row.height = Cm(0.45)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # 음수 단가/금액도 자동 할인 강조
    eff_amt = item.effective_amount()
    if not is_discount:
        if (eff_amt is not None and eff_amt < 0) or (
            item.unit_price is not None and item.unit_price < 0
        ):
            is_discount = True

    # 컬럼 2: 상세 구분 (이름 + 보조 + 종량제 하위) — sub_items 만 더 작은 폰트로
    name_cell = table.cell(row_idx, 2)
    _vcenter(name_cell)
    if is_discount:
        _set_cell_bg(name_cell, DISCOUNT_BG)
    name_p = name_cell.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(0)
    # 할인 강조용 색상/배경 (할인 행일 때만 사용)
    disc_color = DISCOUNT_COLOR if is_discount else None
    disc_bg = DISCOUNT_BG if is_discount else None

    nr = name_p.add_run(item.name)
    _apply_font(nr, font, size_pt=7.5,
                bold=is_discount, color=disc_color)
    if item.name_detail:
        for ln in item.name_detail.split("\n"):
            p_detail = name_cell.add_paragraph()
            p_detail.paragraph_format.space_before = Pt(0)
            p_detail.paragraph_format.space_after = Pt(0)
            r_detail = p_detail.add_run(ln)
            _apply_font(r_detail, font, size_pt=7.5, color=disc_color)
    if item.sub_items:
        sub_lines = []
        for s in item.sub_items:
            if s.label:
                sub_lines.append(f"· {s.label}  {s.spec}")
            elif sub_lines:
                sub_lines[-1] += f"  {s.spec}"
            else:
                sub_lines.append(f"   {s.spec}")
        for ln in sub_lines:
            p_sub = name_cell.add_paragraph()
            p_sub.paragraph_format.space_before = Pt(0)
            p_sub.paragraph_format.space_after = Pt(0)
            r_sub = p_sub.add_run(ln)
            _apply_font(r_sub, font, size_pt=6.5, color=disc_color)

    # 컬럼 3: 기간
    _fill_cell(table.cell(row_idx, 3), item.billing_period or "-",
               font=font, size_pt=7.5, align=WD_ALIGN_PARAGRAPH.CENTER,
               bold=is_discount, color=disc_color, bg=disc_bg)

    # 컬럼 4: 단가
    if item.unit_price is not None:
        up_text = _money(item.unit_price)
        up_align = WD_ALIGN_PARAGRAPH.RIGHT
    elif item.unit_price_text:
        up_text = item.unit_price_text
        up_align = WD_ALIGN_PARAGRAPH.LEFT
    else:
        up_text = "-"
        up_align = WD_ALIGN_PARAGRAPH.CENTER
    _fill_cell(table.cell(row_idx, 4), up_text,
               font=font, size_pt=7.5, align=up_align,
               bold=is_discount, color=disc_color, bg=disc_bg)

    # 컬럼 5: 할인 (할인금액 > 할인율). '할인' 의미상 음수 — '-' 부호 없이 금액만
    if item.discount_amount and item.discount_amount > 0:
        disc_text = _money(item.discount_amount)
    elif item.discount_rate:
        disc_text = f"{int(round(item.discount_rate * 100))}%"
    else:
        disc_text = "-"
    item_has_discount = bool(
        (item.discount_amount and item.discount_amount > 0) or item.discount_rate
    )
    _fill_cell(table.cell(row_idx, 5), disc_text,
               font=font, size_pt=7.5, align=WD_ALIGN_PARAGRAPH.CENTER,
               bold=item_has_discount or is_discount,
               color=DISCOUNT_COLOR if (item_has_discount or is_discount) else None,
               bg=DISCOUNT_BG if (item_has_discount or is_discount) else None)

    # 컬럼 6: 금액
    eff = item.effective_amount()
    if item.amount_text:
        amt_text = item.amount_text
    elif eff is not None:
        amt_text = _money(eff)
    else:
        amt_text = "-"
    _fill_cell(table.cell(row_idx, 6), amt_text,
               font=font, size_pt=8, bold=True,
               align=WD_ALIGN_PARAGRAPH.RIGHT,
               color=disc_color, bg=disc_bg)

    # 컬럼 7: 비고
    _fill_cell(table.cell(row_idx, 7), item.notes or "-",
               font=font, size_pt=7.5, align=WD_ALIGN_PARAGRAPH.LEFT,
               bold=is_discount, color=disc_color, bg=disc_bg)


def _render_unit_notice(doc, document: MembershipQuoteDocument, brand: Brand) -> None:
    font = brand.branding.font_family
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"아래와 같이 견적하오니 참조하시기 바랍니다.    {document.unit_notice}")
    _apply_font(r, font, size_pt=7.5, color=RGBColor(0x77, 0x77, 0x77))


def _render_grand_total(doc, scenario: MembershipScenario,
                        brand: Brand, vat_rate: float = 0.10,
                        total_discount_rate: float | None = None) -> None:
    """시나리오 전체의 공급가액·부가세·합계 금액. 기간별 breakdown 유지."""
    font = brand.branding.font_family
    primary = _hex_to_rgb(brand.branding.colors.primary)
    gross_by_period = scenario_grand_total_by_period(scenario)
    if not gross_by_period:
        return

    def _inline(period_map: dict) -> str:
        return "  /  ".join(f"{p}: {_money(v)}" for p, v in period_map.items())

    td = total_discount_rate or 0
    by_period = {p: v * (1 - td) for p, v in gross_by_period.items()}
    vat_by_period = {p: round(v * vat_rate) for p, v in by_period.items()}
    total_by_period = {p: by_period[p] + vat_by_period[p] for p in by_period}

    pct = int(vat_rate * 100)
    rows = []
    if td > 0:
        td_by_period = {p: gross_by_period[p] - by_period[p] for p in gross_by_period}
        rows.append(("공급가액 (할인 전)", _inline(gross_by_period), False))
        rows.append((f"일괄 할인 ({int(round(td * 100))}%)",
                     _inline({p: -v for p, v in td_by_period.items()}), False))
    rows.extend([
        ("공급가액", _inline(by_period), False),
        (f"부가세 ({pct}%)", _inline(vat_by_period), False),
        ("합계 금액", _inline(total_by_period), True),
    ])

    for idx, (label, value, highlight) in enumerate(rows):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(4 if idx == 0 else 0)
        p.paragraph_format.space_after = Pt(0)
        r_label = p.add_run(f"{label}  ")
        _apply_font(r_label, font, size_pt=7, bold=True,
                    color=primary if highlight else None)
        r_value = p.add_run(value)
        _apply_font(r_value, font,
                    size_pt=8 if highlight else 7,
                    bold=True,
                    color=primary if highlight else None)


def _render_remarks(doc, document: MembershipQuoteDocument, brand: Brand) -> None:
    if not document.remarks:
        return
    font = brand.branding.font_family
    _add_paragraph(doc, "※ 기타 안내", font=font, size_pt=8, bold=True,
                   space_before_pt=4, space_after_pt=0)
    for i, line in enumerate(document.remarks, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.2)
        r = p.add_run(f"{i}. {line}")
        _apply_font(r, font, size_pt=7.5)


def _render_date(doc, document: MembershipQuoteDocument, brand: Brand) -> None:
    font = brand.branding.font_family
    _add_paragraph(doc, document.issued_date.strftime("%Y년 %m월 %d일"),
                   font=font, size_pt=9,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before_pt=4, space_after_pt=2)


def _add_page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _add_tiny_spacer(doc) -> None:
    """섹션 사이 아주 작은 간격 (2pt)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("")
    r.font.size = Pt(2)


# ─── 메인 진입점 ───────────────────────────────────────────

def render_membership_docx(brand: Brand, document: MembershipQuoteDocument,
                           project_root: Path, output_path: Path) -> Path:
    labels = load_labels(project_root)
    vat_rate = labels.quote.vat_rate

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.9)
        section.bottom_margin = Cm(0.9)
        section.left_margin = Cm(0.75)
        section.right_margin = Cm(0.75)

    # QR 견적서와 동일한 흐름: 로고 → 헤더(회사+발급정보) → 수신처(제휴사) →
    # 단위 안내 → 섹션 표 → 합계 → 기타 안내 → 서명(날짜+회사명)
    for s_idx, scenario in enumerate(document.scenarios):
        if s_idx > 0:
            _add_page_break(doc)
        _render_logo(doc, brand, project_root)
        _render_header_qr_style(doc, document, brand)
        _render_counterparty_qr_style(doc, document, brand)
        _render_unit_notice(doc, document, brand)
        for sec_i, sec in enumerate(scenario.sections):
            _render_section_table(doc, sec, brand)
            if sec_i < len(scenario.sections) - 1:
                _add_tiny_spacer(doc)
        if scenario.show_grand_total:
            _render_grand_total(doc, scenario, brand, vat_rate,
                                total_discount_rate=document.total_discount_rate)
        _render_remarks(doc, document, brand)
        _render_signature_qr_style(doc, document, brand)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
